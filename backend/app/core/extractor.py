from tika import parser  # Dùng để trích xuất nội dung từ file PDF
from docx import Document  # Dùng để đọc file Word (.docx)
from pathlib import Path
from io import BytesIO  # Dùng để xử lý file ở bộ nhớ

def extract_text_from_file(file) -> str:
    """
    Hàm gọi chung cho các file định dạng khác nhau.
    """
    if isinstance(file, BytesIO):  # Kiểm tra xem file có phải là BytesIO không
        # Chuyển đổi memoryview thành byte
        file_content = file.getbuffer().tobytes()
        
        # Nếu file là byte, phân loại và xử lý
        if file_content.startswith(b"%PDF"):  # Đặc trưng của file PDF
            return extract_text_from_pdf(file)
        elif file_content.startswith(b"PK"):  # Đặc trưng của file DOCX
            return extract_text_from_word(file)
        elif file_content.startswith(b"\xEF\xBB\xBF"):  # Đặc trưng của file TXT
            return extract_text_from_txt(file)
        else:
            raise ValueError("Unsupported file format in stream.")
    else:
        raise ValueError("Unsupported file format. Only stream files are allowed.")

def extract_text_from_pdf(file) -> str:
    """
    Trích xuất văn bản từ file PDF.
    """
    parsed = parser.from_buffer(file)
    return parsed["content"].strip() if parsed["content"] else ""

def extract_text_from_word(file) -> str:
    """
    Trích xuất văn bản từ file Word (.docx).
    """
    try:
        doc = Document(file)  # Đọc file DOCX với python-docx
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return content
    except Exception as e:
        raise RuntimeError(f"Failed to read Word file: {str(e)}")


def extract_text_from_txt(file) -> str:
    """
    Trích xuất văn bản từ file TXT (hỗ trợ cả đối tượng BytesIO).
    """
    if isinstance(file, BytesIO):  # Nếu là BytesIO, đọc trực tiếp từ bộ nhớ
        return file.read().decode("utf-8").strip()
    else:
        with open(file, "r", encoding="utf-8") as f:  # Nếu là đường dẫn, đọc từ file
            return f.read().strip()

def read_word_file(file) -> str:
    """
    Đọc nội dung từ file Word (.docx) và trả về chuỗi text.
    """
    try:
        doc = Document(file)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return content
    except Exception as e:
        raise RuntimeError(f"Failed to read Word file: {str(e)}")