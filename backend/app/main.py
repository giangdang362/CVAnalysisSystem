from fastapi import FastAPI, Depends, UploadFile, File, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, text, asc
from sqlalchemy.orm import sessionmaker, Session
from typing import List
from datetime import datetime
from app.models import CV, JD  # Assuming models are in app.models
from dotenv import load_dotenv
from docx import Document  # To handle docx files
# from app.invoke_bedrock import invoke_api
from app.invoke_gemini import invoke_gemini_api
import uuid, os, boto3
from PyPDF2 import PdfReader
from pydantic import BaseModel
from typing import List
from botocore.exceptions import ClientError
import logging
from app.models.schemas import CvCreateSchema, CvSchema, JdCreateSchema, JdSchema

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Load environment variables
load_dotenv()

# Database connection settings
DATABASE_URL = os.getenv("DATABASE_URL")

# S3 config
BUCKET_NAME = 'lts-4-aisayhi'
s3_client = boto3.client('s3')

# Initialize database engine and session
engine = create_engine(DATABASE_URL, echo=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

def get_db():
    """Dependency to get database session"""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Initialize FastAPI app
app = FastAPI(
    title="CV and JD Analysis System",
    description="System for analyzing CVs and JDs with AI",
    version="1.0.0",
)

# Add middleware for CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# File upload folder
PROMPT_FILE_OLD = "app/configs/prompt_guide.docx"
PROMPT_FILE= "app/configs/prompt.docx"

# Default route
@app.get("/")
async def root():
    return {"message": "Welcome to CV and JD Analysis System"}

from sqlalchemy.sql import text

@app.get("/roles", tags=["Role Management"])
async def get_roles(db: Session = Depends(get_db)):
    """Get a list of all roles from the enum 'role_type'"""
    try:
        # Truy vấn lấy danh sách các giá trị enum từ database
        query = text("SELECT unnest(enum_range(NULL::role_type)) AS role;")
        result = db.execute(query).fetchall()
        role_list = [row[0] for row in result]

        return {"message": "success", "data": role_list, "count": len(role_list)}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/cv/{id}", tags=["CV Management"])
async def get_cv_by_id(id: int, db: Session = Depends(get_db)):
    """
    Get details of a specific CV by ID.
    """
    try:
        cv = db.query(CV).filter(CV.id == id).first()
        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")
        return {"message": "success", "data": cv}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/jd/{id}", tags=["JD Management"])
async def get_jd_by_id(id: int, db: Session = Depends(get_db)):
    """
    Get details of a specific JD by ID.
    """
    try:
        jd = db.query(JD).filter(JD.id == id).first()
        if not jd:
            raise HTTPException(status_code=404, detail="JD not found")
        return {"message": "success", "data": jd}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/cv/create", tags=["CV Management"], response_model=CvSchema)
async def add_cv(
    payload: CvCreateSchema = Body(...),
    db: Session = Depends(get_db),
):
    try:

        # Thêm CV record vào database
        new_cv = CV(
            name=payload.name,
            path_file=payload.path_file,  # Lưu đường dẫn S3
            expect_salary=payload.expect_salary,
            role=payload.role,
            education=payload.education,
            recruiter=payload.recruiter,
            experience_summary=payload.experience_summary,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
        )
        db.add(new_cv)
        db.commit()
        db.refresh(new_cv)

        return new_cv
    except Exception as e:
        logger.error(f"Error while creating CV: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while creating CV."
        )
@app.put("/cv/{cv_id}", tags=["CV Management"], response_model=CvSchema)
async def update_cv(
    cv_id: int,
    payload: CvCreateSchema = Body(...),  # Dữ liệu cần cập nhật
    db: Session = Depends(get_db),
):
    try:
        # Tìm CV trong cơ sở dữ liệu
        cv = db.query(CV).filter(CV.id == cv_id).first()
        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")

        # Cập nhật các trường
        cv.name = payload.name
        cv.path_file = payload.path_file
        cv.expect_salary = payload.expect_salary
        cv.role = payload.role
        cv.education = payload.education
        cv.recruiter = payload.recruiter
        cv.experience_summary = payload.experience_summary
        cv.updated_at = datetime.utcnow()

        # Lưu thay đổi vào cơ sở dữ liệu
        db.commit()
        db.refresh(cv)

        return cv
    except Exception as e:
        logger.error(f"Error while updating CV: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while updating CV."
        )
@app.delete("/cv/{cv_id}", tags=["CV Management"])
async def delete_cv(cv_id: int, db: Session = Depends(get_db)):
    try:
        # Tìm CV trong cơ sở dữ liệu
        cv = db.query(CV).filter(CV.id == cv_id).first()
        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")

        # Xóa CV khỏi cơ sở dữ liệu
        db.delete(cv)
        db.commit()

        return {"status": "success", "message": f"CV with id {cv_id} has been deleted."}
    except Exception as e:
        logger.error(f"Error while deleting CV: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while deleting CV."
        )
@app.get("/cv", tags=["CV Management"])
async def get_all_cvs(db: Session = Depends(get_db)):
    """Get a list of all CVs"""
    cvs = db.query(CV).order_by(asc(CV.created_at)).all()
    return {"message": "success", "data": cvs, "count": len(cvs)}

@app.post("/jd/create", tags=["JD Management"], response_model=JdSchema)
async def add_jd(
    payload: JdCreateSchema = Body(...),
    db: Session = Depends(get_db),
):
    try:
        # Thêm JD record vào database
        new_jd = JD(
            name=payload.name,
            path_file=payload.path_file,  # Lưu đường dẫn S3
            company_name=payload.company_name,
            role=payload.role,
            level=payload.level,
            languages=payload.languages,
            technical_skill=payload.technical_skill,
            requirement=payload.requirement,
            description=payload.description,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
        )
        db.add(new_jd)
        db.commit()
        db.refresh(new_jd)

        return new_jd
    except Exception as e:
        logger.error(f"Error while creating JD: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while creating JD."
        )
@app.put("/jd/{jd_id}", tags=["JD Management"], response_model=JdSchema)
async def update_jd(
    jd_id: int,
    payload: JdCreateSchema = Body(...),  # Dữ liệu cần cập nhật
    db: Session = Depends(get_db),
):
    try:
        # Tìm JD trong cơ sở dữ liệu
        jd = db.query(JD).filter(JD.id == jd_id).first()
        if not jd:
            raise HTTPException(status_code=404, detail="JD not found")

        # Cập nhật các trường
        jd.name = payload.name
        jd.path_file = payload.path_file
        jd.company_name = payload.company_name
        jd.role = payload.role
        jd.level = payload.level
        jd.languages = payload.languages
        jd.technical_skill = payload.technical_skill
        jd.requirement = payload.requirement
        jd.description = payload.description
        jd.updated_at = datetime.utcnow()

        # Lưu thay đổi vào cơ sở dữ liệu
        db.commit()
        db.refresh(jd)

        return jd
    except Exception as e:
        logger.error(f"Error while updating JD: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while updating JD."
        )
@app.delete("/jd/{jd_id}", tags=["JD Management"])
async def delete_jd(jd_id: int, db: Session = Depends(get_db)):
    try:
        # Tìm JD trong cơ sở dữ liệu
        jd = db.query(JD).filter(JD.id == jd_id).first()
        if not jd:
            raise HTTPException(status_code=404, detail="JD not found")

        # Xóa JD khỏi cơ sở dữ liệu
        db.delete(jd)
        db.commit()

        return {"status": "success", "message": f"JD with id {jd_id} has been deleted."}
    except Exception as e:
        logger.error(f"Error while deleting JD: {str(e)}")  # Log lỗi debug
        raise HTTPException(
            status_code=500, detail="An error occurred while deleting JD."
        )
@app.get("/jd", tags=["JD Management"])
async def get_all_jds(db: Session = Depends(get_db)):
    """Get a list of all JDs"""
    jds = db.query(JD).order_by(asc(JD.created_at)).all()
    return {"message": "success", "data": jds, "count": len(jds)}

@app.post("/matching/cv-to-jds", tags=["Matching"])
async def match_cv_to_jds(cv_id: int, db: Session = Depends(get_db)):
    """Match a CV with JDs based on the role"""
    # Fetch the CV by id
    cv = db.query(CV).filter(CV.id == cv_id).first()
    if not cv:
        return {"message": "error", "detail": "CV not found"}

    # Find JDs that match the CV's role
    matching_jds = db.query(JD).filter(JD.role == cv.role).all()
    ids = [item.id for item in matching_jds]
    return {"message": "success", "data": matching_jds, "ids": ids, "count": len(matching_jds)}

@app.post("/matching/jd-to-cvs", tags=["Matching"])
async def match_jd_to_cvs(jd_id: int, db: Session = Depends(get_db)):
    """Match a JD with CVs based on the role"""
    # Fetch the JD by id
    jd = db.query(JD).filter(JD.id == jd_id).first()
    if not jd:
        return {"message": "error", "detail": "JD not found"}

    # Find CVs that match the JD's role
    matching_cvs = db.query(CV).filter(CV.role == jd.role).all()
    ids = [item.id for item in matching_cvs]
    return {"message": "success", "data": matching_cvs, "ids": ids, "count": len(matching_cvs)}

class RankRequest(BaseModel):
    cv_id: int
    jd_ids: List[int]
@app.post("/matching/cv-to-jds/rank", tags=["Matching"])
async def rank_cv_against_jds(payload: RankRequest, db: Session = Depends(get_db)):
    try:
        # Lấy CV từ DB
        cv = db.query(CV).filter(CV.id == payload.cv_id).first()
        if not cv:
            raise HTTPException(status_code=404, detail="CV not found")

        # Lấy danh sách JD từ DB
        jds = db.query(JD).filter(JD.id.in_(payload.jd_ids)).all()
        if not jds:
            raise HTTPException(status_code=404, detail="No JDs found")

        # Đọc nội dung CV từ S3
        cv_text = ""
        try:
            if cv.path_file.lower().endswith(".pdf"):
                cv_text = read_pdf_from_s3(cv.path_file)
            elif cv.path_file.lower().endswith(".docx"):
                cv_text = read_docx_from_s3(cv.path_file)
            else:
                raise HTTPException(status_code=400, detail="Unsupported CV file format")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error reading CV: {e}")

        # Đọc nội dung Prompt từ file DOCX
        prompt_text = read_docx(PROMPT_FILE_OLD)

        results = []

        # Đọc nội dung JD từ S3 và thực hiện phân tích
        for jd in jds:
            jd_text = ""
            try:
                if jd.path_file.lower().endswith(".pdf"):
                    jd_text = read_pdf_from_s3(jd.path_file)
                elif jd.path_file.lower().endswith(".docx"):
                    jd_text = read_docx_from_s3(jd.path_file)
                else:
                    raise HTTPException(status_code=400, detail="Unsupported JD file format")
            except Exception as e:
                # Ghi nhận lỗi cho JD cụ thể và tiếp tục
                results.append({"name": jd.name, "error": f"Error reading JD: {e}"})
                continue

            # Chuẩn bị prompt
            prompt = f"""
            You are tasked with evaluating a CV against a Job Description (JD) based on the following criteria: Tech Stack, experience, language, and leadership.
            {prompt_text}

            CV:
            {cv_text}

            JD:
            {jd_text}
            """

            # Gọi API phân tích
            response = invoke_gemini_api(prompt)
            
            if "error" in response:
                results.append({"name": jd.name, "error": response["error"]})
                continue

            # Xây dựng kết quả
            result = {
                "name": jd.name,
                "overall_score": response.get("overall_score", 0),
                "tech_stack": response.get("tech_stack", 0),
                "experience": response.get("experience", 0),
                "language": response.get("language", 0),
                "leadership": response.get("leadership", 0),
            }
            results.append(result)

        # Sắp xếp kết quả hợp lệ theo overall_score (giảm dần)
        valid_results = [res for res in results if "overall_score" in res]
        valid_results.sort(key=lambda x: x["overall_score"], reverse=True)

        return {"message": "success", "data": valid_results, "count": len(valid_results), "errors": [res for res in results if "error" in res]}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class RankJdToCvsRequest(BaseModel):
    jd_id: int
    cv_ids: List[int]
@app.post("/matching/jd-to-cvs/rank", tags=["Matching"])
async def rank_jd_against_cvs(
    payload: RankJdToCvsRequest, db: Session = Depends(get_db)
):
    """
    Rank a JD against a list of CVs based on scoring criteria.
    """
    try:
        # Fetch JD details
        jd = db.query(JD).filter(JD.id == payload.jd_id).first()
        if not jd:
            raise HTTPException(status_code=404, detail="JD not found")

        # Fetch CVs from DB
        cvs = db.query(CV).filter(CV.id.in_(payload.cv_ids)).all()
        if not cvs:
            raise HTTPException(status_code=404, detail="No CVs found with the provided IDs")

        # Read JD file from S3
        jd_text = ""
        if jd.path_file.lower().endswith(".pdf"):
            jd_text = read_pdf_from_s3(jd.path_file)
        elif jd.path_file.lower().endswith(".docx"):
            jd_text = read_docx_from_s3(jd.path_file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type for JD")

        # Read the Prompt file
        prompt_text = read_docx(PROMPT_FILE)
        results = []

        # Process each CV
        for cv in cvs:
            # Read CV file from S3
            cv_text = ""
            if cv.path_file.lower().endswith(".pdf"):
                cv_text = read_pdf_from_s3(cv.path_file)
            elif cv.path_file.lower().endswith(".docx"):
                cv_text = read_docx_from_s3(cv.path_file)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type for CV")

            # Prepare prompt input
            prompt = f"""
            You are tasked with evaluating a Job Description (JD) against a CV based on 4 criteria: 
            Tech Stack, experience, language, and leadership.

            {prompt_text}

            JD: 
            {jd_text}

            CV:
            {cv_text}
            """

            # Call Claude API (Gemini API)
            response = invoke_gemini_api(prompt)
            
            if "error" in response:
                raise HTTPException(status_code=500, detail=response["error"])

            # Build result
            result = {
                "name": cv.name,
                "overall_score": response.get("overall_score", 0),
                "tech_stack": response.get("tech_stack", 0),
                "experience": response.get("experience", 0),
                "language": response.get("language", 0),
                "leadership": response.get("leadership", 0),
            }
            results.append(result)

        # Sort results by overall_score
        results.sort(key=lambda x: x["overall_score"], reverse=True)

        return {"message": "success", "data": results, "count": len(results)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload/{path}", tags=["Upload File to S3"])
async def upload_file_s3(
    path: str,
    file: UploadFile = File(...)
    ):
    # Tạo unique filename
    file_name = f"{uuid.uuid4()}-{file.filename}"
    s3_path = f"{path}/{file_name}"
    
    try:
        # Upload file lên S3
        s3_client.upload_fileobj(file.file, BUCKET_NAME, s3_path)
        # Trả lại path để lưu vào DB
        return {"s3_path": s3_path, "message": "File uploaded successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading file: {str(e)}")

# Hàm đọc file trực tiếp từ S3
def read_file_from_s3(path_file: str):
    try:
        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=path_file)
        return response['Body'].read()  # Trả về nội dung file
    except ClientError as e:
        raise HTTPException(status_code=500, detail=f"Error reading file from S3: {e}")

def download_file_from_s3(s3_path):
    # Tạo đường dẫn file tạm
    file_name = os.path.basename(s3_path)
    local_path = f"/tmp/{file_name}"
    
    try:
        # Download file từ S3
        s3_client.download_file(BUCKET_NAME, s3_path, local_path)
        return local_path
    except Exception as e:
        print(f"Error downloading file: {e}")
        return None
    

# Hàm đọc PDF từ S3
def read_pdf_from_s3(path_file: str):
    from PyPDF2 import PdfReader
    from io import BytesIO  # Import BytesIO để chuyển đổi bytes thành file-like object

    # Đọc nội dung file từ S3
    file_content = read_file_from_s3(path_file)

    # Bọc nội dung file bytes vào BytesIO
    pdf_reader = PdfReader(BytesIO(file_content))
    
    # Trích xuất text từ các trang PDF
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()  # PyPDF2 có thể trả về None nếu không đọc được text
    return text


# Hàm đọc DOCX từ S3
def read_docx_from_s3(path_file: str):
    from io import BytesIO
    from docx import Document
    file_content = read_file_from_s3(path_file)
    doc = Document(BytesIO(file_content))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

# Helper function to read docx file
def read_docx(file_path: str) -> str:
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File does not exist: {file_path}")
        content = Document(file_path)
        text = "\n".join([para.text for para in content.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read docx file: {e}")

# Helper function to read pdf file
def read_pdf(file_path: str) -> str:
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File does not exist: {file_path}")
        reader = PdfReader(file_path)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        return text
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read PDF file: {e}")

# Run the application (only for testing purposes, not in production)
if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
